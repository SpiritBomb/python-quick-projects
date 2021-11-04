from docx import Document
from docx.shared import Inches # allows you to resize an image

import pyttsx3 # imports python text to speech

def speak(text):
    pyttsx3.speak(text)

document = Document()    # creating document

# import a picture into document
document.add_picture(
    'me.jpg', 
    width = Inches(2.0)
) 

# name - phone number -email details
name = input("What is your name? ")
speak("Hello " + name + " how are you today?")

speak("What is your phone number?")
phone_number = input("What is your phone number? ")
email = input("What is your email? ")

# data inside document
document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email
)

#about me heading
document.add_heading('About me')
about_me = input("Tell me about yourself ")
document.add_paragraph(about_me)

# work experience heading
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company name ')
from_date = input('From Date ') # to and from date for experience
to_date = input('To Date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True  

experience_details = input(
    'Describe you experience at ' + company + ' '
)
p.add_run(experience_details)


# more experiences
while True:
    has_more_experiences = input("Do you have more experiences? Yes or No ")
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company name ')
        from_date = input('From Date ') # to and from date for experience
        to_date = input('To Date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True  

        experience_details = input(
            'Describe you experience at ' + company + ' '
        )
        p.add_run(experience_details)
    else:
        break  # takes you out of the loop
        

# list of skills
document.add_heading('Skills')
skill = input('Enter your skill: ')

p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more skills Yes or No ? ')
    if has_more_skills.lower() == 'yes':
        
        skill = input('Enter your skill: ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break    

# document footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text= " CV generated using Python programming"

document.save('cv.docx') # saving document name

# run the code below in terminal if (pyttsx3 and python-docx)libraries are not found
# pip3 install -r requirements.txt 
